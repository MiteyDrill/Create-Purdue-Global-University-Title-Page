import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

### Make sure to run {pip install -r requirements.txt} for the docx import.
### Install Pip the python package manager if you haven't already.
### The docx library will be saved in your system-wide or virtual environment's site-packages directory

# Create a new Document object
doc = Document()

# Get the default paragraph style for the document
default_style = doc.styles['Normal']

# Set the line spacing to double-spaced and spacing after to 0pt
default_style.font.name = 'Times New Roman'
default_style.font.size = Pt(12)
default_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
default_style.paragraph_format.space_after = Pt(0)

######

doc.add_paragraph("")  # This adds an empty line (newline).
doc.add_paragraph("")  # This adds an empty line (newline).
doc.add_paragraph("")  # This adds an empty line (newline).
doc.add_paragraph("")  # This adds an empty line (newline).

# APA Title Page
title_page = doc.sections[0]
title_page.start_type


# Title (Formatted as Title Case and Bold)
title_text = input("Title: ")
title_paragraph = doc.add_paragraph()
title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_paragraph.add_run(title_text.title())  # Convert to title case
title_run.bold = True
title_run.font.size = Pt(12)

doc.add_paragraph("")  # This adds an empty line (newline).

# Author(s)
author = input("Your Name: ")
author_paragraph = doc.add_paragraph(author)
author_run = author_paragraph.runs[0]
author_run.font.name = 'Times New Roman'
author_run.font.size = Pt(12)
author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Institutional Affiliation
affiliation = "Purdue Global University"
affiliation_paragraph = doc.add_paragraph(affiliation)
affiliation_run = affiliation_paragraph.runs[0]
affiliation_run.font.name = 'Times New Roman'
affiliation_run.font.size = Pt(12)
affiliation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Class Name:
af1 = input("Class i.e (CM220: Digital Rhetorics): ")

affiliation = af1
affiliation_paragraph = doc.add_paragraph(affiliation)
affiliation_run = affiliation_paragraph.runs[0]
affiliation_run.font.name = 'Times New Roman'
affiliation_run.font.size = Pt(12)
affiliation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Professor Name
professor = input("Professor Name: ")
professor_paragraph = doc.add_paragraph(professor)
professor_run = professor_paragraph.runs[0]
professor_run.font.name = 'Times New Roman'
professor_run.font.size = Pt(12)
professor_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date = input("Date i.e (September 2, 2023): ")
date_paragraph = doc.add_paragraph(date)
date_run = date_paragraph.runs[0]
date_run.font.name = 'Times New Roman'
date_run.font.size = Pt(12)
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Content Page Break
doc.add_page_break()

# Content Section
content_section = doc.add_paragraph("Content")
content_section.runs[0].font.size = Pt(12)

# Content Page Break
doc.add_page_break()

# References Section
references_section = doc.add_paragraph("References")
references_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
references_section.runs[0].bold = True
references_section.runs[0].font.size = Pt(12)

# Set paragraph indentation to hanging
references_section.paragraph_format.first_line_indent = Pt(-12)

######

document_name = input("File Name (e.g., CS220_Ben_CompetencyAssessment_Part1): ")

# Prompt the user for a custom file path
custom_file_path = input("Enter the custom file path and filename (e.g., C:/xampp3/htdocs/PythonScripts):")

try:
    # Normalize the path for the operating system
    custom_file_path = os.path.normpath(custom_file_path)

    # Save the document to the custom file path
    doc.save(custom_file_path + "/" + document_name + ".docx")

    print(f"Document saved to: {custom_file_path}")
except Exception as e:
    print(f"An error occurred while saving the document: {e}")